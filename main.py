import sys
import json
import zipfile
import argparse
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from io import BytesIO
from lxml import etree

class PrecisionEditor:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        # This map acts as our "RAM" to look up objects by the IDs we generate
        self.id_map = {}
        self._generate_map()

    def _generate_map(self):
        """
        Parses the document to build the structured JSON and assigns unique IDs.
        We rebuild this every time we ask for a map to ensure IDs match current state.
        """
        self.id_map = {}
        structure = {
            "sections": [{"id": "s1", "headings": []}],
            "tables": [],
            "metadata": {"total_paragraphs": 0, "total_tables": 0}
        }

        current_section = structure["sections"][0]
        current_heading = None

        # We need a root 'container' for content before the first heading
        root_content = {
            "id": "h_root",
            "level": 0,
            "text": "Root",
            "paragraphs": []
        }
        current_heading = root_content
        current_section["headings"].append(root_content)

        p_count = 0
        r_count = 0

        for i, p in enumerate(self.doc.paragraphs):
            p_id = f"p{i}"
            self.id_map[p_id] = p
            p_count += 1

            # Detect if it's a heading based on style name
            style_name = "Normal" # Default fallback
            if p.style is not None and p.style.name:
                style_name = p.style.name

            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split(' ')[-1])
                except:
                    level = 1

                new_heading = {
                    "id": p_id,
                    "level": level,
                    "text": p.text,
                    "paragraphs": []
                }
                current_section["headings"].append(new_heading)
                current_heading = new_heading
            else:
                # It's standard content
                p_data = {
                    "id": p_id,
                    "text": p.text[:50] + "..." if len(p.text) > 50 else p.text,
                    "runs": []
                }

                # Map Runs (the granular text chunks with formatting)
                for j, run in enumerate(p.runs):
                    r_id = f"{p_id}_r{j}"
                    self.id_map[r_id] = run
                    p_data["runs"].append({
                        "id": r_id,
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic
                    })

                current_heading["paragraphs"].append(p_data)

        # Map Tables (simplified for this prototype)
        for i, table in enumerate(self.doc.tables):
            t_id = f"t{i}"
            self.id_map[t_id] = table
            structure["tables"].append({"id": t_id, "rows": len(table.rows)})

        structure["metadata"]["total_paragraphs"] = p_count
        structure["metadata"]["total_tables"] = len(self.doc.tables)
        self.structure_cache = structure
        return structure

    def get_map_json(self):
        # Refresh map to account for edits
        data = self._generate_map()
        return json.dumps(data, indent=2)

    def replace_text(self, element_id, new_text):
        """
        Surgical replacement. If it's a Run, we preserve style.
        """
        target = self.id_map.get(element_id)
        if not target:
            return f"Error: ID {element_id} not found."

        # Check if it's a Run or a Paragraph
        if hasattr(target, 'font'): # It's a Run
            # modifying run.text keeps the formatting (bold/italic/etc) intact
            target.text = new_text
            return f"Updated Run {element_id}. Formatting preserved."

        elif isinstance(target, Paragraph):
            # If replacing a whole paragraph, we usually lose run-level formatting
            # Strategy: clear content and add new text to first run?
            # For safety in this specific challenge, let's just update the text
            target.clear()
            target.add_run(new_text)
            return f"Updated Paragraph {element_id}. Note: Complex inner formatting reset."

    def insert_after(self, element_id, text):
        """
        The Hard Part: Inserting after a specific element requires OXML manipulation.
        python-docx only has 'insert_paragraph_before', not 'after'.
        """
        target = self.id_map.get(element_id)
        if not target:
            return f"Error: ID {element_id} not found."

        if isinstance(target, Paragraph):
            # We need to access the XML parent to insert a sibling
            # Logic: Create new paragraph, move it in XML tree
            new_p = self.doc.add_paragraph(text)
            # Remove from bottom of doc (where add_paragraph puts it)
            # And insert it specifically after our target

            # Access underlying XML element (_p)
            target_xml = target._element
            new_p_xml = new_p._element

            # XML Insertion: parent.insert(index, element)
            parent = target_xml.getparent()
            parent.insert(parent.index(target_xml) + 1, new_p_xml)

            # Try to inherit style if possible
            new_p.style = target.style

            return f"Inserted new paragraph after {element_id}."
        else:
            return "Error: INSERT_AFTER currently only supported for Paragraphs (p IDs)."

    def delete_element(self, element_id):
        target = self.id_map.get(element_id)
        if not target:
            return f"Error: {element_id} not found."

        if isinstance(target, Paragraph):
            p = target._element
            p.getparent().remove(p)
            p._p = p._element = None
            return f"Deleted {element_id}"

        # If it's a run, simpler
        if hasattr(target, 'font'):
            # Emptying text is safer than XML removal for runs to avoid corruption
            target.text = ""
            return f"Cleared text from Run {element_id}"

    def format_element(self, element_id, prop, value):
        target = self.id_map.get(element_id)
        if not target:
            return "ID not found."

        # Normalize value
        val_bool = str(value).lower() == 'true'

        if prop == 'bold':
            if hasattr(target, 'font'): target.font.bold = val_bool
            else: target.runs[0].font.bold = val_bool
        elif prop == 'italic':
            if hasattr(target, 'font'): target.font.italic = val_bool
            else: target.runs[0].font.italic = val_bool
        elif prop == 'size':
            # Value expected as integer pt
            pt_val = int(value)
            if hasattr(target, 'font'): target.font.size = Pt(pt_val)
            else: target.style.font.size = Pt(pt_val)

        return f"Formatted {element_id}: {prop}={value}"

    def save(self, output_path):
        self.doc.save(output_path)
        return f"Saved to {output_path}"

    def validate(self, output_path):
        """
        Integrity Check:
        1. Is it a valid Zip?
        2. Can we parse the internal document.xml?
        """
        try:
            if not zipfile.is_zipfile(output_path):
                return "FAIL: File is not a valid zip container."

            with zipfile.ZipFile(output_path, 'r') as z:
                # Try to parse the main XML body
                xml_content = z.read('word/document.xml')
                try:
                    etree.fromstring(xml_content)
                except etree.XMLSyntaxError:
                    return "FAIL: Internal XML is corrupt/malformed."

            return "PASS: Document structure and XML are valid."
        except Exception as e:
            return f"FAIL: Validation error: {str(e)}"

# --- CLI INTERFACE ---

def main():
    print("--- Precision Document Editor Prototype ---")
    print("Type 'help' for commands or 'exit' to quit.")

    editor = None

    while True:
        try:
            user_input = input("> ").strip().split(" ")
            cmd = user_input[0].lower()
            args = user_input[1:]

            if cmd == "exit":
                break

            elif cmd == "help":
                print("\nCommands:")
                print("  load [filename]              - Load a .docx file")
                print("  map                          - Show document structure JSON")
                print("  replace [id] [text...]       - Replace text in ID")
                print("  insert_after [id] [text...]  - Insert paragraph after ID")
                print("  delete [id]                  - Delete element")
                print("  format [id] [prop] [val]     - Set prop (bold/italic) to true/false")
                print("  save [filename]              - Save output")
                print("  validate [filename]          - Check integrity\n")

            elif cmd == "load":
                if not args: print("Usage: load [filename]"); continue
                try:
                    editor = PrecisionEditor(args[0])
                    print(f"Loaded {args[0]}")
                    print(f"Stats: {editor.structure_cache['metadata']}")
                except Exception as e:
                    print(f"Error loading: {e}")

            elif cmd == "map":
                if not editor: print("No document loaded."); continue
                print(editor.get_map_json())

            elif cmd == "replace":
                if not editor: print("No document loaded."); continue
                if len(args) < 2: print("Usage: replace [id] [new text]"); continue
                tgt_id = args[0]
                new_text = " ".join(args[1:])
                # Handle quoted strings simply
                new_text = new_text.strip('"').strip("'")
                print(editor.replace_text(tgt_id, new_text))

            elif cmd == "insert_after":
                if not editor: print("No document loaded."); continue
                if len(args) < 2: print("Usage: insert_after [id] [new text]"); continue
                tgt_id = args[0]
                new_text = " ".join(args[1:])
                new_text = new_text.strip('"').strip("'")
                print(editor.insert_after(tgt_id, new_text))

            elif cmd == "delete":
                if not editor: print("No document loaded."); continue
                if len(args) < 1: print("Usage: delete [id]"); continue
                print(editor.delete_element(args[0]))

            elif cmd == "format":
                if not editor: print("No document loaded."); continue
                if len(args) < 3: print("Usage: format [id] [prop] [value]"); continue
                print(editor.format_element(args[0], args[1], args[2]))

            elif cmd == "save":
                if not editor: print("No document loaded."); continue
                if len(args) < 1: print("Usage: save [filename]"); continue
                print(editor.save(args[0]))

            elif cmd == "validate":
                if len(args) < 1: print("Usage: validate [filename]"); continue
                print(editor.validate(args[0]))

            else:
                print("Unknown command.")

        except Exception as e:
            print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()