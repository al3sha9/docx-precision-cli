import json
import zipfile
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree


class DocxEngine:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.id_map = {}
        self.structure_cache = {}
        self._generate_map()

    def _generate_map(self):
        """
        Parses the document to build the structured JSON and assigns unique IDs.
        We rebuild this every time we ask for a map to ensure IDs match current state.
        """
        self.id_map = {}
        structure = {
            "sections": [{"id": "s1", "headings": []}],
            "tables": [],
            "metadata": {"total_paragraphs": 0, "total_tables": 0}
        }

        current_section = structure["sections"][0]
        current_heading = None

        root_content = {
            "id": "h_root",
            "level": 0,
            "text": "Root",
            "paragraphs": []
        }
        current_heading = root_content
        current_section["headings"].append(root_content)

        p_count = 0

        for i, p in enumerate(self.doc.paragraphs):
            p_id = f"p{i}"
            self.id_map[p_id] = p
            p_count += 1

            style_name = "Normal"
            if p.style is not None and p.style.name:
                style_name = p.style.name

            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split(' ')[-1])
                except Exception:
                    level = 1

                new_heading = {
                    "id": p_id,
                    "level": level,
                    "text": p.text,
                    "paragraphs": []
                }
                current_section["headings"].append(new_heading)
                current_heading = new_heading
            else:
                p_data = {
                    "id": p_id,
                    "text": p.text[:50] + "..." if len(p.text) > 50 else p.text,
                    "runs": []
                }

                for j, run in enumerate(p.runs):
                    r_id = f"{p_id}_r{j}"
                    self.id_map[r_id] = run
                    p_data["runs"].append({
                        "id": r_id,
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic
                    })

                current_heading["paragraphs"].append(p_data)

        for i, table in enumerate(self.doc.tables):
            t_id = f"t{i}"
            self.id_map[t_id] = table
            structure["tables"].append({"id": t_id, "rows": len(table.rows)})

        structure["metadata"]["total_paragraphs"] = p_count
        structure["metadata"]["total_tables"] = len(self.doc.tables)
        self.structure_cache = structure
        return structure

    def get_structure_data(self) -> list[dict]:
        """
        Traverse the document and return a flat list of dictionaries.
        Each dict represents a paragraph or table with keys:
          - id: unique identifier (e.g., 'p0', 't0')
          - text: string content
          - style: style name
          - type: 'text' or 'table'
        Also rebuilds id_map for other methods.
        """
        self.id_map = {}
        result: list[dict] = []

        for i, p in enumerate(self.doc.paragraphs):
            p_id = f"p{i}"
            self.id_map[p_id] = p

            style_name = "Normal"
            if p.style is not None and p.style.name:
                style_name = p.style.name

            # Register runs in id_map
            for j, run in enumerate(p.runs):
                r_id = f"{p_id}_r{j}"
                self.id_map[r_id] = run

            result.append({
                "id": p_id,
                "text": p.text,
                "style": style_name,
                "type": "text"
            })

        for i, table in enumerate(self.doc.tables):
            t_id = f"t{i}"
            self.id_map[t_id] = table
            result.append({
                "id": t_id,
                "text": f"[Table with {len(table.rows)} rows]",
                "style": "",
                "type": "table"
            })

        return result

    def get_structure_json(self) -> str:
        """
        Return get_structure_data() serialized as a JSON string.
        """
        data = self.get_structure_data()
        return json.dumps(data, indent=2)

    def map_structure(self):
        """
        Return a human-readable structured map as a Python dict.
        Internally uses get_structure_data for id_map population,
        then formats for CLI display.
        """
        flat_data = self.get_structure_data()

        # Build CLI-friendly hierarchical structure
        structure = {
            "sections": [{"id": "s1", "headings": []}],
            "tables": [],
            "metadata": {"total_paragraphs": 0, "total_tables": 0}
        }
        current_section = structure["sections"][0]
        root_content = {
            "id": "h_root",
            "level": 0,
            "text": "Root",
            "paragraphs": []
        }
        current_section["headings"].append(root_content)
        current_heading = root_content

        p_count = 0
        for item in flat_data:
            if item["type"] == "text":
                p_count += 1
                style_name = item["style"]
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.split(" ")[-1])
                    except Exception:
                        level = 1
                    new_heading = {
                        "id": item["id"],
                        "level": level,
                        "text": item["text"],
                        "paragraphs": []
                    }
                    current_section["headings"].append(new_heading)
                    current_heading = new_heading
                else:
                    p_obj = self.id_map.get(item["id"])
                    runs_list = []
                    if p_obj:
                        for j, run in enumerate(p_obj.runs):
                            r_id = f"{item['id']}_r{j}"
                            runs_list.append({
                                "id": r_id,
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic
                            })
                    p_data = {
                        "id": item["id"],
                        "text": item["text"][:50] + "..." if len(item["text"]) > 50 else item["text"],
                        "runs": runs_list
                    }
                    current_heading["paragraphs"].append(p_data)
            elif item["type"] == "table":
                table_obj = self.id_map.get(item["id"])
                rows = len(table_obj.rows) if table_obj else 0
                structure["tables"].append({"id": item["id"], "rows": rows})

        structure["metadata"]["total_paragraphs"] = p_count
        structure["metadata"]["total_tables"] = len(structure["tables"])
        self.structure_cache = structure
        return structure

    def get_map_json(self):
        """Return the current structured map as a JSON string."""
        data = self._generate_map()
        return json.dumps(data, indent=2)

    def replace_text(self, element_id, new_text):
        """
        Surgical replacement. If it's a Run, we preserve style.
        Returns a status message string.
        """
        target = self.id_map.get(element_id)
        if not target:
            return f"Error: ID {element_id} not found."

        if hasattr(target, 'font'):
            target.text = new_text
            return f"Updated Run {element_id}. Formatting preserved."

        elif isinstance(target, Paragraph):
            try:
                target.clear()
            except Exception:
                # Fallback: remove existing runs by setting text to empty
                for r in list(target.runs):
                    r.text = ""
            target.add_run(new_text)
            return f"Updated Paragraph {element_id}. Note: Complex inner formatting reset."

        return f"Error: Unsupported element type for {element_id}"

    def insert_after(self, element_id, text):
        """
        Insert a new paragraph after the specified paragraph ID.
        Returns a status message string.
        """
        target = self.id_map.get(element_id)
        if not target:
            return f"Error: ID {element_id} not found."

        if isinstance(target, Paragraph):
            new_p = self.doc.add_paragraph(text)
            target_xml = target._element
            new_p_xml = new_p._element
            parent = target_xml.getparent()
            parent.insert(parent.index(target_xml) + 1, new_p_xml)
            new_p.style = target.style
            return f"Inserted new paragraph after {element_id}."
        else:
            return "Error: INSERT_AFTER currently only supported for Paragraphs (p IDs)."

    def delete_element(self, element_id):
        """Delete the referenced element. Returns a status message string."""
        target = self.id_map.get(element_id)
        if not target:
            return f"Error: {element_id} not found."

        if isinstance(target, Paragraph):
            p = target._element
            p.getparent().remove(p)
            p._p = p._element = None
            return f"Deleted {element_id}"

        if hasattr(target, 'font'):
            target.text = ""
            return f"Cleared text from Run {element_id}"

        return f"Error: Unsupported element type for {element_id}"

    def format_element(self, element_id, prop, value):
        """Format the element; returns a status message string."""
        target = self.id_map.get(element_id)
        if not target:
            return "ID not found."

        val_bool = str(value).lower() == 'true'

        if prop == 'bold':
            if hasattr(target, 'font'):
                target.font.bold = val_bool
            else:
                if target.runs:
                    target.runs[0].font.bold = val_bool
        elif prop == 'italic':
            if hasattr(target, 'font'):
                target.font.italic = val_bool
            else:
                if target.runs:
                    target.runs[0].font.italic = val_bool
        elif prop == 'size':
            pt_val = int(value)
            if hasattr(target, 'font'):
                target.font.size = Pt(pt_val)
            else:
                target.style.font.size = Pt(pt_val)

        return f"Formatted {element_id}: {prop}={value}"

    def save(self, output_path):
        """Save the current document. Returns a status message string."""
        self.doc.save(output_path)
        return f"Saved to {output_path}"

    # ------------------------------------------------------------------ #
    #  API Methods for External/Agent Usage                              #
    # ------------------------------------------------------------------ #

    def read_chunk(self, start_index: int, size: int) -> str:
        """
        Return a JSON string containing a slice of paragraphs.

        This method is useful for processing large documents in chunks,
        allowing an external caller to paginate through the document.

        Args:
            start_index (int): Zero-based index of the first paragraph to include.
            size (int): Maximum number of paragraphs to return.

        Returns:
            str: A JSON-encoded array of paragraph dictionaries. Each dict
                 contains keys: id, text, style, type.

        Raises:
            ValueError: If start_index is negative or size is not positive.

        Example:
            >>> engine.read_chunk(0, 10)
            '[{"id": "p0", "text": "Hello", "style": "Normal", "type": "text"}, ...]'
        """
        if start_index < 0:
            raise ValueError("start_index must be >= 0")
        if size <= 0:
            raise ValueError("size must be > 0")

        data = self.get_structure_data()
        chunk = data[start_index : start_index + size]
        return json.dumps(chunk, indent=2)

    def insert_translation(
        self, target_id: str, translation_text: str, style: str = None
    ) -> bool:
        """
        Insert a new paragraph immediately after the paragraph identified by target_id.

        This is designed for translation workflows where translated text is
        inserted directly after the source paragraph.

        Args:
            target_id (str): The ID of the existing paragraph (e.g., 'p0', 'p5').
            translation_text (str): The text content for the new paragraph.
            style (str, optional): Style name to apply to the new paragraph.
                                   If None, inherits the style from target_id.

        Returns:
            bool: True if insertion succeeded.

        Raises:
            ValueError: If target_id is not found or does not reference a paragraph.

        Example:
            >>> engine.insert_translation('p3', 'Translated sentence.', style='Normal')
            True
        """
        # Ensure id_map is current
        if not self.id_map:
            self.get_structure_data()

        target = self.id_map.get(target_id)
        if target is None:
            raise ValueError(f"ID '{target_id}' not found in the document.")

        if not isinstance(target, Paragraph):
            raise ValueError(
                f"ID '{target_id}' does not reference a paragraph; "
                "insert_translation only supports paragraph IDs."
            )

        # Create new paragraph at end, then relocate via OXML
        new_p = self.doc.add_paragraph(translation_text)
        target_xml = target._element
        new_p_xml = new_p._element
        parent = target_xml.getparent()
        parent.insert(parent.index(target_xml) + 1, new_p_xml)

        # Apply style
        if style:
            try:
                new_p.style = style
            except KeyError:
                # Style not found; fall back to target style
                new_p.style = target.style
        else:
            new_p.style = target.style

        return True

    def save_file(self, output_path: str) -> str:
        """
        Save the document to the specified path and return the absolute path.

        This is a convenience wrapper intended for external API callers who
        need the resolved filesystem path.

        Args:
            output_path (str): Destination file path (relative or absolute).

        Returns:
            str: The absolute path of the saved file.

        Raises:
            OSError: If the file cannot be written (e.g., permission denied).

        Example:
            >>> engine.save_file('output/translated.docx')
            '/home/user/project/output/translated.docx'
        """
        import os

        self.doc.save(output_path)
        return os.path.abspath(output_path)

    @staticmethod
    def validate(output_path):
        """
        Integrity Check:
        1. Is it a valid Zip?
        2. Can we parse the internal document.xml?
        Returns a PASS/FAIL message string.
        """
        try:
            if not zipfile.is_zipfile(output_path):
                return "FAIL: File is not a valid zip container."

            with zipfile.ZipFile(output_path, 'r') as z:
                xml_content = z.read('word/document.xml')
                try:
                    etree.fromstring(xml_content)
                except etree.XMLSyntaxError:
                    return "FAIL: Internal XML is corrupt/malformed."

            return "PASS: Document structure and XML are valid."
        except Exception as e:
            return f"FAIL: Validation error: {str(e)}"
